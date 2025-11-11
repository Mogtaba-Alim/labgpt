"""
API route - JSON endpoints for AJAX requests.
"""

from flask import Blueprint, jsonify
from unified_app.services.state_manager import StateManager
from unified_app.utils.log_parser import LogParser
from unified_app.models import Job
from unified_app.extensions import db

api_bp = Blueprint('api', __name__)


@api_bp.route('/project/<int:project_id>/status')
def project_status(project_id):
    """
    Get comprehensive project status including completion status and active jobs.

    Returns:
        JSON with project info, completion flags, active jobs, and navigation permissions
    """
    status = StateManager.get_project_status(project_id)
    if not status:
        return jsonify({'error': 'Project not found'}), 404
    return jsonify(status)


@api_bp.route('/job/<int:job_id>/status')
def job_status(job_id):
    """
    Get job status with progress information parsed from logs.

    Returns:
        JSON with job status, progress percentage, current step, and errors
    """
    job = Job.query.get(job_id)
    if not job:
        return jsonify({'error': 'Job not found'}), 404

    # Parse log file for latest progress
    if job.log_file and job.status in ['running', 'pending']:
        log_info = LogParser.parse_log_file(job.log_file, job.job_type)
        # Update job progress in database
        job.progress_percentage = log_info['progress_percentage']
        job.current_step = log_info['current_step']

        # If log shows errors and job is still running, mark as failed
        if log_info.get('errors') and job.status == 'running':
            job.status = 'failed'
            job.error_message = '; '.join(log_info['errors'][:3])  # First 3 errors

        db.session.commit()

    return jsonify(job.to_dict())


@api_bp.route('/job/<int:job_id>/logs')
def job_logs(job_id):
    """
    Get recent log lines for a job (last 50 lines).

    Returns:
        JSON with job_id and list of log lines
    """
    job = Job.query.get(job_id)
    if not job or not job.log_file:
        return jsonify({'error': 'Job or log file not found'}), 404

    log_lines = LogParser.tail_log_file(job.log_file, num_lines=50)

    return jsonify({
        'job_id': job_id,
        'log_lines': log_lines
    })


@api_bp.route('/training/<int:project_id>/metrics')
def training_metrics(project_id):
    """
    Get training metrics (loss, steps, epochs) for real-time display.

    Returns:
        JSON with current loss, step info, and loss history
    """
    # Get active training job
    active_jobs = StateManager.get_active_jobs(project_id, job_type='training')

    if not active_jobs:
        return jsonify({'error': 'No active training job found'}), 404

    job = active_jobs[0]

    if not job.log_file:
        return jsonify({'error': 'No log file available'}), 404

    # Parse training metrics from log
    metrics = LogParser.get_training_metrics(job.log_file)

    return jsonify({
        'job_id': job.id,
        'status': job.status,
        'metrics': metrics
    })


@api_bp.route('/chat/message', methods=['POST'])
def chat_message_default():
    """
    Process chat message without a project (default mode).

    Uses:
    - Default LabGPT model from HuggingFace
    - Default RAG index (if available)

    Request JSON:
        - messages: List of message dicts with 'role' and 'content'
        - use_rag: Boolean, whether to use RAG (default: True)
        - project_id: Optional project ID to use project-specific model

    Returns:
        JSON with:
            - response: Generated text
            - citations: List of citation dicts (if RAG used)
            - context_used: Boolean
            - model_used: String indicating which model was used
    """
    from flask import request
    from pathlib import Path
    from unified_app.models import Project, ChatMessage
    from unified_app.services.inference_adapter import InferenceAdapter
    from unified_app.config import Config
    import logging

    try:
        # Get request data
        data = request.get_json()
        messages = data.get('messages', [])
        use_rag = data.get('use_rag', True)  # Default True for default mode
        project_id = data.get('project_id')  # Optional project ID

        if not messages:
            return jsonify({'error': 'No messages provided'}), 400

        # Determine which mode to use
        if project_id:
            # Project mode - use project-specific model
            project = Project.query.get(project_id)
            if not project:
                return jsonify({'error': f'Project {project_id} not found'}), 404

            project_dir = Config.PROJECTS_BASE_DIR / project.project_dir
            model_path = project_dir / 'training_outputs' / 'final_model'

            if not model_path.exists():
                return jsonify({'error': 'Project model not found. Using default model instead.', 'fallback': True}), 404

            # RAG index from project
            rag_index_dir = None
            if use_rag and project.rag_completed:
                rag_index_dir = str(project_dir / 'rag_index')

            model_used = f"Project: {project.name}"
            model_path_str = str(model_path)
        else:
            # Default mode - use HuggingFace model
            model_path_str = None  # InferenceAdapter will use default
            model_used = f"Default LabGPT ({Config.DEFAULT_MODEL_ADAPTER})"

            # Check for default RAG index
            rag_index_dir = None
            if use_rag:
                default_rag = Config.DEFAULT_RAG_INDEX_DIR
                if default_rag.exists() and (default_rag / 'documents_metadata.json').exists():
                    rag_index_dir = str(default_rag)

        # Initialize inference adapter
        adapter = InferenceAdapter(
            model_path=model_path_str,
            rag_index_dir=rag_index_dir
        )

        # Add system message if not present
        if not messages or messages[0].get('role') != 'system':
            from inference import LABGPT_SYSTEM
            messages = [{'role': 'system', 'content': LABGPT_SYSTEM}] + messages

        # Generate response
        result = adapter.chat(
            messages=messages,
            use_rag=use_rag,
            max_new_tokens=600,
            temperature=0.4,
            top_p=0.9
        )

        # Add metadata about which model was used
        result['model_used'] = model_used
        result['rag_available'] = rag_index_dir is not None

        # Save message to database (only if project specified)
        if project_id:
            user_message = messages[-1]['content'] if messages[-1]['role'] == 'user' else ''
            chat_message = ChatMessage(
                project_id=project_id,
                user_message=user_message,
                assistant_response=result['response'],
                used_rag=use_rag
            )
            db.session.add(chat_message)
            db.session.commit()

        return jsonify(result)

    except Exception as e:
        logging.error(f"Error generating chat response: {e}")
        return jsonify({'error': str(e)}), 500


@api_bp.route('/projects/list', methods=['GET'])
def projects_list():
    """
    Get list of all projects with their availability status.

    Returns:
        JSON with:
            - projects: List of project dicts with {id, name, has_model, has_rag}
    """
    from unified_app.models import Project
    from unified_app.config import Config

    projects = Project.query.all()
    project_list = []

    for project in projects:
        project_dir = Config.PROJECTS_BASE_DIR / project.project_dir
        model_path = project_dir / 'training_outputs' / 'final_model'
        rag_index_path = project_dir / 'rag_index'

        project_list.append({
            'id': project.id,
            'name': project.name,
            'has_model': model_path.exists(),
            'has_rag': rag_index_path.exists() and project.rag_completed
        })

    return jsonify({'projects': project_list})


@api_bp.route('/projects/<int:project_id>/chat-status', methods=['GET'])
def project_chat_status(project_id):
    """
    Get project status for chat interface.

    Returns:
        JSON with:
            - model_available: Boolean
            - model_path: String if available
            - rag_available: Boolean
            - rag_path: String if available
            - project_name: String
    """
    from unified_app.models import Project
    from unified_app.config import Config

    project = Project.query.get_or_404(project_id)
    project_dir = Config.PROJECTS_BASE_DIR / project.project_dir

    model_path = project_dir / 'training_outputs' / 'final_model'
    rag_index_path = project_dir / 'rag_index'

    return jsonify({
        'project_name': project.name,
        'model_available': model_path.exists(),
        'model_path': str(model_path) if model_path.exists() else None,
        'rag_available': rag_index_path.exists() and project.rag_completed,
        'rag_path': str(rag_index_path) if rag_index_path.exists() else None
    })


@api_bp.route('/chat/check-rag-status', methods=['GET'])
def check_rag_status():
    """
    Check if default RAG index is available.

    Returns:
        JSON with:
            - has_default: Boolean, whether default index exists
            - default_path: String, path to default index if exists
            - is_valid: Boolean, whether index has required files
    """
    from unified_app.config import Config

    default_rag = Config.DEFAULT_RAG_INDEX_DIR

    has_default = default_rag.exists()
    is_valid = False

    if has_default:
        # Check for required files
        required_files = ['documents_metadata.json', 'faiss.index', 'chunks.json']
        is_valid = all((default_rag / f).exists() for f in required_files)

    return jsonify({
        'has_default': has_default,
        'default_path': str(default_rag) if has_default else None,
        'is_valid': is_valid
    })


@api_bp.route('/chat/<int:project_id>/message', methods=['POST'])
def chat_message(project_id):
    """
    Process chat message and return response.

    Request JSON:
        - messages: List of message dicts with 'role' and 'content'
        - use_rag: Boolean, whether to use RAG

    Returns:
        JSON with:
            - response: Generated text
            - citations: List of citation dicts (if RAG used)
            - context_used: Boolean
    """
    from flask import request
    from pathlib import Path
    from unified_app.models import Project, ChatMessage
    from unified_app.services.inference_adapter import InferenceAdapter
    from unified_app.config import Config
    import logging

    project = Project.query.get_or_404(project_id)
    project_dir = Config.PROJECTS_BASE_DIR / project.project_dir

    try:
        # Get request data
        data = request.get_json()
        messages = data.get('messages', [])
        use_rag = data.get('use_rag', False)

        if not messages:
            return jsonify({'error': 'No messages provided'}), 400

        # Get model path from training outputs
        model_path = project_dir / 'training_outputs' / 'final_model'
        if not model_path.exists():
            return jsonify({'error': 'Trained model not found. Please complete training first.'}), 404

        # Get RAG index directory if RAG enabled
        rag_index_dir = None
        if use_rag and project.rag_completed:
            rag_index_dir = str(project_dir / 'rag_index')

        # Initialize inference adapter
        adapter = InferenceAdapter(
            model_path=str(model_path),
            rag_index_dir=rag_index_dir
        )

        # Add system message if not present
        if not messages or messages[0].get('role') != 'system':
            from inference import LABGPT_SYSTEM
            messages = [{'role': 'system', 'content': LABGPT_SYSTEM}] + messages

        # Generate response
        result = adapter.chat(
            messages=messages,
            use_rag=use_rag,
            max_new_tokens=600,
            temperature=0.4,
            top_p=0.9
        )

        # Save message to database
        user_message = messages[-1]['content'] if messages[-1]['role'] == 'user' else ''
        chat_message = ChatMessage(
            project_id=project_id,
            user_message=user_message,
            assistant_response=result['response'],
            used_rag=use_rag
        )
        db.session.add(chat_message)
        db.session.commit()

        return jsonify(result)

    except Exception as e:
        logging.error(f"Error generating chat response: {e}")
        return jsonify({'error': str(e)}), 500


@api_bp.route('/chat/<int:project_id>/history', methods=['GET'])
def chat_history(project_id):
    """
    Get chat history for this project.

    Returns:
        JSON with messages list
    """
    from unified_app.models import Project, ChatMessage

    project = Project.query.get_or_404(project_id)

    try:
        # Get recent messages (last 50)
        messages = ChatMessage.query.filter_by(project_id=project_id)\
            .order_by(ChatMessage.created_at.asc())\
            .limit(50)\
            .all()

        # Convert to message format
        message_list = []
        for msg in messages:
            message_list.append({
                'role': 'user',
                'content': msg.user_message
            })
            message_list.append({
                'role': 'assistant',
                'content': msg.assistant_response,
                'citations': []  # Could store citations in DB if needed
            })

        return jsonify({'messages': message_list})

    except Exception as e:
        import logging
        logging.error(f"Error fetching chat history: {e}")
        return jsonify({'error': str(e)}), 500


@api_bp.route('/chat/<int:project_id>/clear', methods=['POST'])
def chat_clear(project_id):
    """
    Clear chat history for this project.

    Returns:
        JSON with success status
    """
    from unified_app.models import Project, ChatMessage

    project = Project.query.get_or_404(project_id)

    try:
        # Delete all messages for this project
        ChatMessage.query.filter_by(project_id=project_id).delete()
        db.session.commit()

        return jsonify({'success': True})

    except Exception as e:
        import logging
        logging.error(f"Error clearing chat history: {e}")
        db.session.rollback()
        return jsonify({'error': str(e)}), 500


@api_bp.route('/grant/<int:project_id>/generate', methods=['POST'])
def grant_generate(project_id):
    """
    Generate content for a grant section.

    Request JSON:
        - section: Section name

    Returns:
        JSON with:
            - content: Generated section text
            - citations: List of citation dicts
            - quality_feedback: Quality assessment
    """
    from flask import request
    from pathlib import Path
    from unified_app.models import Project
    from unified_app.services.inference_adapter import InferenceAdapter
    from unified_app.config import Config
    import logging

    project = Project.query.get_or_404(project_id)
    project_dir = Config.PROJECTS_BASE_DIR / project.project_dir

    try:
        data = request.get_json()
        section = data.get('section', '')

        if not section:
            return jsonify({'error': 'Section name required'}), 400

        # Section prompts (simplified versions)
        SECTION_PROMPTS = {
            'Background': "Write a comprehensive Background section that establishes the scientific foundation for this research. Include current state of knowledge, knowledge gaps, and research justification.",
            'Specific Aims': "Write clear Specific Aims that articulate what you will accomplish. Include 2-3 numbered aims with rationale, approach, and expected outcomes.",
            'Significance': "Write a Significance section explaining why this research matters and its potential impact on the field.",
            'Innovation': "Write an Innovation section highlighting novel approaches and unique aspects of your work.",
            'Approach': "Write an Approach section detailing your research methods and experimental design.",
            'Environment': "Write an Environment section describing available resources and institutional support.",
            'Bibliography': "Compile all cited references in proper format based on the citations used in previous sections."
        }

        prompt = SECTION_PROMPTS.get(section, f"Write the {section} section for this research grant proposal.")

        # Get model and RAG paths
        model_path = project_dir / 'training_outputs' / 'final_model'
        rag_index_dir = str(project_dir / 'rag_index') if project.rag_completed else None

        if not model_path.exists():
            return jsonify({'error': 'Trained model not found'}), 404

        # Initialize inference adapter
        adapter = InferenceAdapter(
            model_path=str(model_path),
            rag_index_dir=rag_index_dir
        )

        # Generate with RAG
        result = adapter.generate(
            query=f"{prompt}\n\nWrite a well-structured, academic {section} section.",
            use_rag=True,
            top_k=5,
            max_new_tokens=800,
            temperature=0.4,
            top_p=0.9
        )

        # Simple quality feedback
        word_count = len(result['response'].split())
        quality_feedback = f"Generated {word_count} words. "
        if result['citations']:
            quality_feedback += f"Included {len(result['citations'])} citations."
        else:
            quality_feedback += "No citations found - consider enabling RAG."

        return jsonify({
            'content': result['response'],
            'citations': result['citations'],
            'quality_feedback': quality_feedback
        })

    except Exception as e:
        logging.error(f"Error generating grant section: {e}")
        return jsonify({'error': str(e)}), 500


@api_bp.route('/grant/<int:project_id>/refine', methods=['POST'])
def grant_refine(project_id):
    """
    Refine existing grant section content.

    Request JSON:
        - section: Section name
        - current_draft: Current section content

    Returns:
        JSON with refined content and citations
    """
    from flask import request
    from pathlib import Path
    from unified_app.models import Project
    from unified_app.services.inference_adapter import InferenceAdapter
    from unified_app.config import Config
    import logging

    project = Project.query.get_or_404(project_id)
    project_dir = Config.PROJECTS_BASE_DIR / project.project_dir

    try:
        data = request.get_json()
        section = data.get('section', '')
        current_draft = data.get('current_draft', '')

        if not section or not current_draft:
            return jsonify({'error': 'Section and current_draft required'}), 400

        # Get model and RAG paths
        model_path = project_dir / 'training_outputs' / 'final_model'
        rag_index_dir = str(project_dir / 'rag_index') if project.rag_completed else None

        # Initialize inference adapter
        adapter = InferenceAdapter(
            model_path=str(model_path),
            rag_index_dir=rag_index_dir
        )

        # Refine prompt
        refine_prompt = f"Improve and refine the following {section} section:\n\n{current_draft}\n\nProvide a refined version with better clarity, stronger arguments, and additional relevant citations."

        result = adapter.generate(
            query=refine_prompt,
            use_rag=True,
            top_k=5,
            max_new_tokens=800,
            temperature=0.4,
            top_p=0.9
        )

        word_count = len(result['response'].split())
        quality_feedback = f"Refined version: {word_count} words with {len(result['citations'])} citations."

        return jsonify({
            'content': result['response'],
            'citations': result['citations'],
            'quality_feedback': quality_feedback
        })

    except Exception as e:
        logging.error(f"Error refining grant section: {e}")
        return jsonify({'error': str(e)}), 500


@api_bp.route('/grant/<int:project_id>/save', methods=['POST'])
def grant_save(project_id):
    """
    Save grant draft to database.

    Request JSON:
        - drafts: Dict of {section_name: {content, citations, ...}}

    Returns:
        JSON with success status
    """
    from flask import request
    from unified_app.models import Project, SectionDraft
    import json

    project = Project.query.get_or_404(project_id)

    try:
        data = request.get_json()
        drafts = data.get('drafts', {})

        for section_name, section_data in drafts.items():
            # Check if draft exists
            existing_draft = SectionDraft.query.filter_by(
                project_id=project_id,
                section_name=section_name
            ).first()

            if existing_draft:
                # Update existing
                existing_draft.content = section_data.get('content', '')
                existing_draft.citations_json = json.dumps(section_data.get('citations', []))
                existing_draft.version += 1
            else:
                # Create new
                new_draft = SectionDraft(
                    project_id=project_id,
                    section_name=section_name,
                    content=section_data.get('content', ''),
                    citations_json=json.dumps(section_data.get('citations', [])),
                    version=1
                )
                db.session.add(new_draft)

        db.session.commit()
        return jsonify({'success': True})

    except Exception as e:
        import logging
        logging.error(f"Error saving grant draft: {e}")
        db.session.rollback()
        return jsonify({'error': str(e)}), 500


@api_bp.route('/grant/<int:project_id>/draft', methods=['GET'])
def grant_draft(project_id):
    """
    Get saved grant draft.

    Returns:
        JSON with drafts dict
    """
    from unified_app.models import Project, SectionDraft
    import json

    project = Project.query.get_or_404(project_id)

    try:
        drafts = {}
        section_drafts = SectionDraft.query.filter_by(project_id=project_id).all()

        for draft in section_drafts:
            drafts[draft.section_name] = {
                'content': draft.content,
                'citations': json.loads(draft.citations_json) if draft.citations_json else [],
                'version': draft.version
            }

        return jsonify({'drafts': drafts})

    except Exception as e:
        import logging
        logging.error(f"Error loading grant draft: {e}")
        return jsonify({'error': str(e)}), 500


@api_bp.route('/grant/<int:project_id>/export', methods=['POST'])
def grant_export(project_id):
    """
    Export grant as DOCX file.

    Request JSON:
        - drafts: Dict of section drafts
        - format: 'docx' or 'pdf'

    Returns:
        File download
    """
    from flask import request, send_file
    from pathlib import Path
    from unified_app.models import Project
    from unified_app.config import Config
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import io

    project = Project.query.get_or_404(project_id)

    try:
        data = request.get_json()
        drafts = data.get('drafts', {})
        format_type = data.get('format', 'docx')

        # Create DOCX document
        doc = Document()

        # Add title
        title = doc.add_heading(project.name, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add sections in order
        section_order = ['Background', 'Specific Aims', 'Significance', 'Innovation', 'Approach', 'Environment', 'Bibliography']

        for section_name in section_order:
            if section_name in drafts:
                # Add section heading
                doc.add_heading(section_name, 1)

                # Add content
                content = drafts[section_name].get('content', '')
                doc.add_paragraph(content)

                # Add citations
                citations = drafts[section_name].get('citations', [])
                if citations:
                    doc.add_heading('References', 2)
                    for citation in citations:
                        cite_text = f"{citation.get('source', 'Unknown')} - Page {citation.get('page', 'N/A')}"
                        doc.add_paragraph(cite_text, style='List Bullet')

                # Add page break except for last section
                if section_name != section_order[-1]:
                    doc.add_page_break()

        # Save to bytes
        doc_io = io.BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)

        return send_file(
            doc_io,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f'{project.name}_grant.docx'
        )

    except Exception as e:
        import logging
        logging.error(f"Error exporting grant: {e}")
        return jsonify({'error': str(e)}), 500
