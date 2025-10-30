#!/usr/bin/env python3
"""
enhanced_app.py

Modern Grant Generation Web Application (v2.0) with:
- RAG/pipeline.py integration (hybrid retrieval, reranking, query expansion)
- inference.py message format with trained LabGPT model
- Per-project RAG indices with optional base corpus
- Async document indexing with Celery
- SQLite database for projects, drafts, and versioning
- Multi-turn refinement with conversation history
- Consistent generation parameters across all sections
- Quality assessment with citation checks
- Draft version tracking and lineage
"""

import os
import io
import json
import tempfile
from pathlib import Path
from typing import List, Dict
from flask import Flask, request, render_template, redirect, url_for, session, send_file, flash, jsonify
from werkzeug.utils import secure_filename
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
import logging

from corpus_manager import GrantCorpusManager
from database import GrantDatabase
from enhanced_grant_service_v2 import EnhancedGrantService, SectionResult
from tasks import index_documents_async

# Initialize Flask app
app = Flask(__name__)
app.secret_key = os.environ.get('SECRET_KEY', 'your-secret-key-change-in-production')

# Configuration
UPLOAD_FOLDER = 'temp_uploads'
ALLOWED_EXTENSIONS = {'pdf', 'txt', 'docx', 'doc'}
MAX_CONTENT_LENGTH = 16 * 1024 * 1024  # 16MB max file size

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = MAX_CONTENT_LENGTH

# Ensure upload folder exists
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# Initialize services
# Base RAG index from main pipeline (set via environment variable)
base_rag_index = os.environ.get('BASE_RAG_INDEX', None)
corpus_manager = GrantCorpusManager(
    base_index_dir=base_rag_index,
    projects_dir='grant_generation/projects'
)
db = GrantDatabase('grant_generation/grants.db')
grant_service = EnhancedGrantService(corpus_manager, db)

# Grant sections in logical order (matching enhanced_grant_service_v2.py)
SECTIONS_ORDER = [
    "Background",
    "Specific Aims",
    "Significance",
    "Innovation",
    "Approach",
    "Environment",
    "Bibliography"
]

logging.basicConfig(level=logging.INFO)

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def save_uploaded_files(files) -> List[str]:
    """Save uploaded files and return file paths"""
    file_paths = []
    for file in files:
        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(filepath)
            file_paths.append(filepath)
    return file_paths

def get_current_project_id():
    """Get current project ID from Flask session"""
    return session.get('project_id')

@app.route('/', methods=['GET', 'POST'])
def index():
    """Enhanced home page with document upload and project creation"""
    if request.method == 'POST':
        # Get grant details
        grant_title = request.form.get('title', '').strip()
        grant_overview = request.form.get('overview', '').strip()

        if not grant_title or not grant_overview:
            flash('Please provide both grant title and overview.', 'error')
            return render_template('enhanced_index.html')

        try:
            # Create project in database first (generates UUID)
            project_id = db.create_project(
                title=grant_title,
                overview=grant_overview,
                index_dir="",  # Will be updated after creating RAG index
                has_base=(base_rag_index is not None)
            )

            # Create project RAG index with base corpus copy
            project_index_dir = corpus_manager.create_project(
                project_id=project_id,
                copy_base=(base_rag_index is not None)
            )

            # Update project with index directory
            db.update_project(project_id, index_dir=project_index_dir)

            # Store project_id in Flask session
            session['project_id'] = project_id
            session['grant_title'] = grant_title
            session['grant_overview'] = grant_overview

            # Handle file uploads
            uploaded_files = request.files.getlist('supporting_docs')
            file_paths = save_uploaded_files(uploaded_files)

            if file_paths:
                # Trigger async indexing
                task = index_documents_async.apply_async(
                    args=[project_id, file_paths, base_rag_index]
                )

                # Store task ID in session for status polling
                session['indexing_task_id'] = task.id
                session['indexing_in_progress'] = True

                flash(
                    f'Project created! Indexing {len(file_paths)} documents in background. '
                    'You can start working on sections while indexing completes.',
                    'success'
                )
            else:
                flash('Project created successfully!', 'success')

            # Redirect to first section
            return redirect(url_for('section_workflow', section_name=SECTIONS_ORDER[0]))

        except Exception as e:
            flash(f'Error creating project: {str(e)}', 'error')
            logging.error(f"Project creation error: {e}", exc_info=True)
            return render_template('enhanced_index.html')

    return render_template('enhanced_index.html')

@app.route('/indexing_status')
def indexing_status():
    """API endpoint to check async indexing progress"""
    task_id = session.get('indexing_task_id')

    if not task_id:
        return jsonify({
            'status': 'no_task',
            'message': 'No indexing task in progress'
        })

    # Check task status
    from celery.result import AsyncResult
    task = AsyncResult(task_id)

    if task.state == 'PENDING':
        response = {
            'status': 'pending',
            'message': 'Task is waiting to be executed...',
            'current': 0,
            'total': 0
        }
    elif task.state == 'PROCESSING':
        response = {
            'status': 'processing',
            'message': task.info.get('status', 'Processing...'),
            'current': task.info.get('current', 0),
            'total': task.info.get('total', 0)
        }
    elif task.state == 'SUCCESS':
        result = task.result
        response = {
            'status': 'complete',
            'message': 'Indexing complete!',
            'files_processed': result.get('files_processed', 0),
            'total_chunks': result.get('total_chunks', 0),
            'cache_hit_rate': result.get('cache_hit_rate', '0%')
        }
        # Clear task from session
        session.pop('indexing_task_id', None)
        session['indexing_in_progress'] = False
    elif task.state == 'FAILURE':
        response = {
            'status': 'failed',
            'message': f'Indexing failed: {str(task.info)}',
            'error': str(task.info)
        }
        # Clear task from session
        session.pop('indexing_task_id', None)
        session['indexing_in_progress'] = False
    else:
        response = {
            'status': 'unknown',
            'message': f'Unknown task state: {task.state}'
        }

    return jsonify(response)

@app.route('/section/<section_name>', methods=['GET', 'POST'])
def section_workflow(section_name):
    """Enhanced section editing workflow with modern RAG and inference"""
    project_id = get_current_project_id()
    grant_overview = session.get('grant_overview')

    if not project_id or not grant_overview:
        flash('Please start by creating a project.', 'error')
        return redirect(url_for('index'))

    if section_name not in SECTIONS_ORDER:
        flash('Invalid section name.', 'error')
        return redirect(url_for('index'))

    # Get current section index
    current_index = SECTIONS_ORDER.index(section_name)

    # Handle POST request (generation, refinement, or saving)
    if request.method == 'POST':
        action = request.form.get('action')

        if action == 'generate':
            # Generate section content using RAGPipeline + inference.py
            try:
                additional_context = request.form.get('additional_context', '').strip()

                result = grant_service.generate_section(
                    project_id=project_id,
                    section_name=section_name,
                    overview=grant_overview,
                    additional_context=additional_context
                )

                flash(
                    f'Generated {section_name} section (Quality: {result.quality_score:.2f}, '
                    f'{len(result.citations)} citations, v{result.version})',
                    'success'
                )

            except Exception as e:
                flash(f'Error generating section: {str(e)}', 'error')
                logging.error(f"Section generation error: {e}", exc_info=True)

        elif action == 'refine':
            # Refine with user feedback (multi-turn conversation)
            user_feedback = request.form.get('feedback', '').strip()

            if user_feedback:
                try:
                    result = grant_service.refine_section(
                        project_id=project_id,
                        section_name=section_name,
                        feedback=user_feedback
                    )

                    flash(
                        f'Refined {section_name} section (Quality: {result.quality_score:.2f}, v{result.version})',
                        'success'
                    )

                except Exception as e:
                    flash(f'Error refining section: {str(e)}', 'error')
                    logging.error(f"Section refinement error: {e}", exc_info=True)

        elif action == 'next':
            # Move to next section
            if current_index < len(SECTIONS_ORDER) - 1:
                next_section = SECTIONS_ORDER[current_index + 1]
                return redirect(url_for('section_workflow', section_name=next_section))
            else:
                # All sections completed, go to finalization
                return redirect(url_for('finalize'))

    # Get latest draft for current section
    latest_draft = db.get_latest_draft(project_id, section_name)

    # Get project progress
    progress_data = grant_service.get_project_progress(project_id)

    # Calculate overall progress
    completed_count = sum(1 for status in progress_data.values() if status['completed'])
    progress = {
        'current_section': section_name,
        'current_index': current_index,
        'total_sections': len(SECTIONS_ORDER),
        'completed_count': completed_count,
        'percentage': (completed_count / len(SECTIONS_ORDER)) * 100
    }

    # Check if indexing is in progress
    indexing_in_progress = session.get('indexing_in_progress', False)

    return render_template('enhanced_section.html',
                         section_name=section_name,
                         latest_draft=latest_draft,
                         progress=progress,
                         progress_data=progress_data,
                         sections_order=SECTIONS_ORDER,
                         indexing_in_progress=indexing_in_progress)

@app.route('/navigate/<direction>/<section_name>')
def navigate_section(direction, section_name):
    """Navigate between sections"""
    try:
        current_index = SECTIONS_ORDER.index(section_name)
    except ValueError:
        return redirect(url_for('index'))

    if direction == 'next' and current_index < len(SECTIONS_ORDER) - 1:
        next_section = SECTIONS_ORDER[current_index + 1]
        return redirect(url_for('section_workflow', section_name=next_section))
    elif direction == 'prev' and current_index > 0:
        prev_section = SECTIONS_ORDER[current_index - 1]
        return redirect(url_for('section_workflow', section_name=prev_section))
    elif direction == 'overview':
        return redirect(url_for('grant_overview'))

    return redirect(url_for('section_workflow', section_name=section_name))

@app.route('/overview')
def grant_overview():
    """Show grant overview and progress"""
    project_id = get_current_project_id()
    grant_title = session.get('grant_title')
    grant_overview = session.get('grant_overview')

    if not project_id or not grant_overview:
        return redirect(url_for('index'))

    # Get project from database
    project = db.get_project(project_id)

    # Get project progress
    progress_data = grant_service.get_project_progress(project_id)

    # Get uploaded documents
    documents = db.get_project_documents(project_id)

    return render_template('grant_overview.html',
                         project=project,
                         grant_title=grant_title,
                         grant_overview=grant_overview,
                         progress_data=progress_data,
                         documents=documents,
                         sections_order=SECTIONS_ORDER)

@app.route('/finalize')
def finalize():
    """Final review and export preparation"""
    project_id = get_current_project_id()
    grant_title = session.get('grant_title')
    grant_overview = session.get('grant_overview')

    if not project_id or not grant_overview:
        return redirect(url_for('index'))

    # Get project from database
    project = db.get_project(project_id)

    # Get all section drafts
    section_drafts = {}
    for section_name in SECTIONS_ORDER:
        draft = db.get_latest_draft(project_id, section_name)
        if draft:
            section_drafts[section_name] = draft

    # Check if all sections are completed
    missing_sections = [s for s in SECTIONS_ORDER if s not in section_drafts]

    return render_template('enhanced_finalize.html',
                         project=project,
                         grant_title=grant_title,
                         grant_overview=grant_overview,
                         section_drafts=section_drafts,
                         missing_sections=missing_sections,
                         sections_order=SECTIONS_ORDER)

@app.route('/download')
def download_pdf():
    """Generate and download PDF of completed grant"""
    project_id = get_current_project_id()
    grant_title = session.get('grant_title', 'Grant Proposal')
    grant_overview = session.get('grant_overview')

    if not project_id or not grant_overview:
        return redirect(url_for('index'))

    # Get all section drafts
    section_drafts = {}
    for section_name in SECTIONS_ORDER:
        draft = db.get_latest_draft(project_id, section_name)
        if draft:
            section_drafts[section_name] = draft

    # Create PDF
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    story = []
    styles = getSampleStyleSheet()

    # Title
    title = Paragraph(grant_title, styles['Title'])
    story.append(title)
    story.append(Spacer(1, 12))

    # Grant Overview
    overview_title = Paragraph("Grant Overview", styles['Heading1'])
    story.append(overview_title)
    overview_text = Paragraph(grant_overview, styles['Normal'])
    story.append(overview_text)
    story.append(Spacer(1, 12))

    # Sections
    for section_name in SECTIONS_ORDER:
        if section_name in section_drafts:
            draft = section_drafts[section_name]
            section_title = Paragraph(section_name, styles['Heading1'])
            story.append(section_title)
            section_text = Paragraph(draft['content'], styles['Normal'])
            story.append(section_text)
            story.append(Spacer(1, 12))

    doc.build(story)
    buffer.seek(0)

    return send_file(buffer, as_attachment=True, download_name='grant_proposal.pdf', mimetype='application/pdf')

@app.route('/export/docx')
def export_docx():
    """Export grant proposal as DOCX with MLA citations"""
    project_id = get_current_project_id()
    grant_title = session.get('grant_title', 'Grant Proposal')
    grant_overview = session.get('grant_overview')

    if not project_id or not grant_overview:
        flash('No active project to export', 'error')
        return redirect(url_for('index'))

    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

        # Create document
        doc = Document()

        # Title
        title = doc.add_heading(grant_title, level=0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Grant Overview
        doc.add_heading('Grant Overview', level=1)
        doc.add_paragraph(grant_overview)

        # Get all section drafts with citations
        all_citations = []
        citation_map = {}  # Map citation text to MLA format

        for section_name in SECTIONS_ORDER:
            draft = db.get_latest_draft(project_id, section_name)
            if draft:
                # Add section heading
                doc.add_heading(section_name, level=1)

                # Add section content
                doc.add_paragraph(draft['content'])

                # Collect citations
                for citation in draft['citations']:
                    citation_text = citation.get('citation', citation.get('text', ''))
                    if citation_text and citation_text not in citation_map:
                        # Format citation in MLA style
                        mla_citation = format_mla_citation(citation)
                        citation_map[citation_text] = mla_citation
                        all_citations.append(mla_citation)

        # Add Bibliography section with MLA citations
        if all_citations:
            doc.add_page_break()
            doc.add_heading('Works Cited', level=1)

            # Sort citations alphabetically (MLA requirement)
            all_citations.sort()

            for citation in all_citations:
                p = doc.add_paragraph(citation)
                # Hanging indent for MLA format
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.first_line_indent = Inches(-0.5)

        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return send_file(
            buffer,
            as_attachment=True,
            download_name=f'{grant_title.replace(" ", "_")}.docx',
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    except ImportError:
        flash('python-docx is not installed. Please install it: pip install python-docx', 'error')
        return redirect(url_for('finalize'))
    except Exception as e:
        flash(f'Error exporting DOCX: {str(e)}', 'error')
        logging.error(f"DOCX export error: {e}", exc_info=True)
        return redirect(url_for('finalize'))

def format_mla_citation(citation: Dict) -> str:
    """
    Format citation in MLA 9th edition style.

    Args:
        citation: Citation dict with metadata

    Returns:
        MLA-formatted citation string
    """
    # Extract metadata
    text = citation.get('text', '')
    source = citation.get('citation', '')
    score = citation.get('score', 0)

    # Try to parse author, title, publication from source string
    # Common formats: "Author. Title. Publication, Year."
    # or "Title (Year)" or just "filename.pdf"

    # Simple heuristic-based MLA formatting
    if '(' in source and ')' in source:
        # Format: "Title (Year)" or "Author (Year). Title"
        parts = source.split('(')
        title_author = parts[0].strip()
        year = parts[1].split(')')[0].strip() if len(parts) > 1 else ''

        if '.' in title_author:
            # Format: "Author. Title"
            author, title = title_author.split('.', 1)
            return f'{author.strip()}. {title.strip()}. {year}.'
        else:
            # Format: "Title"
            return f'{title_author}. {year}.'

    elif source.endswith('.pdf') or source.endswith('.txt'):
        # Format: filename
        filename = source.replace('.pdf', '').replace('.txt', '').replace('_', ' ')
        return f'"{filename}." Document.'

    else:
        # Generic format
        return f'{source}.'

@app.route('/api/section-quality/<section_name>')
def get_section_quality(section_name):
    """API endpoint to get section quality assessment"""
    project_id = get_current_project_id()
    if not project_id:
        return jsonify({'error': 'No active project'}), 404

    draft = db.get_latest_draft(project_id, section_name)
    if not draft:
        return jsonify({'error': 'No draft found'}), 404

    return jsonify({
        'section_name': section_name,
        'quality_score': draft['quality_score'],
        'quality_feedback': draft['quality_feedback'],
        'version': draft['version'],
        'word_count': len(draft['content'].split()),
        'citation_count': len(draft['citations'])
    })

@app.route('/api/draft-history/<section_name>')
def get_draft_history(section_name):
    """API endpoint to get draft version history"""
    project_id = get_current_project_id()
    if not project_id:
        return jsonify({'error': 'No active project'}), 404

    history = db.get_draft_history(project_id, section_name)
    return jsonify(history)

@app.route('/clear-session')
def clear_session():
    """Clear session and start over"""
    # Note: In production, you may want to keep projects in database
    # and just clear the Flask session to allow switching between projects

    # Clear Flask session
    session.clear()

    # Clean up uploaded files
    try:
        for file in os.listdir(UPLOAD_FOLDER):
            file_path = os.path.join(UPLOAD_FOLDER, file)
            if os.path.isfile(file_path):
                os.remove(file_path)
    except Exception as e:
        logging.warning(f"Error cleaning up upload folder: {e}")

    flash('Session cleared. Starting fresh.', 'info')
    return redirect(url_for('index'))

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000) 